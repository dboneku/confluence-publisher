#!/usr/bin/env python3
"""
Confluence Publisher
Usage: python publish.py [--mapping mapping.csv] [--file path/or/url] [--dry-run]
"""

# ---------------------------------------------------------------------------
# Bootstrap — auto-install dependencies into a venv on first run
# ---------------------------------------------------------------------------
import subprocess
import sys
from pathlib import Path

_VENV = Path(__file__).parent.parent / ".venv"
_SENTINEL = _VENV / ".deps-installed"
_REQS = Path(__file__).parent / "requirements.txt"
_SCRIPT_DIR = Path(__file__).parent

if str(_SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPT_DIR))

def _bootstrap():
    if not _VENV.exists():
        print("First run: creating virtual environment for dependencies...")
        subprocess.check_call([sys.executable, "-m", "venv", str(_VENV)])

    venv_python = _VENV / ("Scripts/python.exe" if sys.platform == "win32" else "bin/python3")

    if not _SENTINEL.exists():
        print("Installing dependencies (one-time setup)...")
        subprocess.check_call([str(venv_python), "-m", "pip", "install", "-q",
                               "-r", str(_REQS)])
        _SENTINEL.touch()

    # Check if we are already running *inside* the venv using sys.prefix —
    # more reliable than comparing executable paths (both may resolve to the same binary).
    already_in_venv = Path(sys.prefix).resolve() == _VENV.resolve()
    if not already_in_venv:
        result = subprocess.run([str(venv_python)] + sys.argv)
        sys.exit(result.returncode)

try:
    import requests          # noqa: F401 — quick check; full import below
    import dotenv            # noqa: F401
except ImportError:
    _bootstrap()

# ---------------------------------------------------------------------------

import os
import re
import json
import base64
import argparse
import requests
from publisherlib.confluence_api import (
    build_space_tree as _build_space_tree,
    create_page as _create_page,
    delete_page as _delete_page,
    fetch_page_adf as _fetch_page_adf,
    find_existing_page as _find_existing_page,
    get_page_version as _get_page_version,
    list_child_pages as _list_child_pages,
    list_spaces as _list_spaces,
    page_url as _page_url,
    resolve_parent_id as _resolve_parent_id,
    resolve_space_id as _resolve_space_id,
    update_page as _update_page,
    walk_descendant_pages as _walk_descendant_pages,
    walk_space_pages as _walk_space_pages,
)
from publisherlib.adf_tools import (
    adf_to_html,
    adf_to_markdown as _adf_to_markdown,
    build_doc_control_footer,
    build_doc_control_header,
    bullet_list,
    diff_adf,
    extract_classification_from_adf as _extract_classification_from_adf,
    extract_doc_id_from_title,
    extract_text_from_adf as _extract_text_from_adf,
    has_doc_control_header as _has_doc_control_header,
    heading,
    info_panel,
    metadata_table,
    node_text as _node_text,
    paragraph,
    strip_doc_control_blocks as _strip_doc_control_blocks,
    text_node,
    to_adf,
    wrap_with_print_controls,
)
from dotenv import load_dotenv
from publisherlib.policy import (
    check_adf_against_style_policy as _check_adf_against_style_policy,
    extract_required_headings_from_policy as _extract_required_headings_from_policy_impl,
    fix_adf_heading_numbers as _fix_adf_heading_numbers,
    inject_regulation_doc_id as _inject_regulation_doc_id,
    strip_title_heading_from_adf as _strip_title_heading_from_adf,
)
from publisherlib.project_config import (
    CONFIG_SCHEMA_VERSION as PROJECT_CONFIG_SCHEMA_VERSION,
    REGULATION_CONFIG_FILE as PROJECT_REGULATION_CONFIG_FILE,
    STYLE_POLICY_FILE as PROJECT_STYLE_POLICY_FILE,
    load_regulation_config,
    load_style_policy,
    normalize_title,
    save_regulation_config,
    save_style_policy,
    warn as _project_warn,
)
from publisherlib.templates import (
    REQUIRED_SECTIONS,
    check_template_sections,
    detect_template_from_text,
    validate_naming_convention,
)

# load_dotenv() with no args calls find_dotenv(), which crashes when the script
# is run from stdin (heredoc). Explicitly load from cwd where .env lives.
load_dotenv(Path(os.getcwd()) / '.env')

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

ATLASSIAN_URL  = os.environ.get("ATLASSIAN_URL", "").rstrip("/")
ATLASSIAN_EMAIL = os.environ.get("ATLASSIAN_EMAIL", "")
ATLASSIAN_TOKEN = os.environ.get("ATLASSIAN_API_TOKEN", "")

# ---------------------------------------------------------------------------
# Regulation catalog
# ---------------------------------------------------------------------------

# ISO 27001:2022 mandatory and recommended document catalog.
# Keys are the organization's document IDs (NN-ISMS format).
# Values are the canonical document names used for fuzzy title matching.
ISO27001_DOCS = {
    "00-ISMS": "Master List of Documents",
    "01-ISMS": "Scope of the ISMS",
    "02-ISMS": "Information Security Policy",
    "03-ISMS": "Information Security Risk Assessment Process",
    "04-ISMS": "Information Security Risk Treatment Plan",
    "05-ISMS": "Statement of Applicability",
    "06-ISMS": "Information Security Objectives",
    "07-ISMS": "Competence and Training Records",
    "08-ISMS": "Monitoring Measurement Analysis and Evaluation",
    "09-ISMS": "Internal Audit Program",
    "10-ISMS": "Internal Audit Results",
    "11-ISMS": "Management Review Results",
    "12-ISMS": "Nonconformities and Corrective Actions",
    "13-ISMS": "Asset Inventory",
    "14-ISMS": "Acceptable Use of Information and Assets Policy",
    "15-ISMS": "Information Classification Policy",
    "16-ISMS": "Access Control Policy",
    "17-ISMS": "Identity Management Policy",
    "18-ISMS": "Supplier Security Policy",
    "19-ISMS": "Incident Response Plan",
    "20-ISMS": "Business Continuity Plan",
    "21-ISMS": "Cryptography Policy",
    "22-ISMS": "Clear Desk and Clear Screen Policy",
    "23-ISMS": "Remote Working Policy",
    "24-ISMS": "Change Management Policy",
    "25-ISMS": "Secure Development Policy",
    "26-ISMS": "Vulnerability Management Policy",
    "27-ISMS": "Network Security Policy",
    "28-ISMS": "Log Management Policy",
    "29-ISMS": "Configuration Management Policy",
    "30-ISMS": "Data Retention and Disposal Policy",
}

REGULATION_CONFIGS = {
    "iso27001": {
        "name":     "ISO/IEC 27001:2022",
        "doc_id_suffix": "ISMS",
        "docs":     ISO27001_DOCS,
    },
}

_REGULATION_CONFIG_FILE = PROJECT_REGULATION_CONFIG_FILE
_CONFIG_SCHEMA_VERSION = PROJECT_CONFIG_SCHEMA_VERSION


def _warn(message: str):
    _project_warn(message)

_STYLE_POLICY_FILE = PROJECT_STYLE_POLICY_FILE


def _page_id_from_url(url: str) -> str | None:
    """Extract numeric Confluence page ID from a URL like .../pages/123456/..."""
    m = re.search(r'/pages/(\d+)', url)
    return m.group(1) if m else None


def _extract_section(text: str, section_heading: str) -> str:
    """Extract a named section from plain text or markdown.

    Finds the first heading line that contains `section_heading` (case-insensitive,
    ignoring leading # chars) and returns everything up to the next heading of
    equal or higher level.  Returns the full text when the section is not found.
    """
    lines = text.splitlines()
    heading_re = re.compile(r'^(#{1,6})\s+(.+)')
    target = section_heading.strip().lower()
    start_idx = None
    start_level = 1

    for i, line in enumerate(lines):
        match = heading_re.match(line)
        text_part = match.group(2).strip().lower() if match else line.strip().lower()
        if target in text_part or text_part in target:
            start_idx = i
            start_level = len(match.group(1)) if match else 1
            break

    if start_idx is None:
        return text

    result = [lines[start_idx]]
    for line in lines[start_idx + 1:]:
        match = heading_re.match(line)
        if match and len(match.group(1)) <= start_level:
            break
        result.append(line)
    return '\n'.join(result)


def _extract_required_headings_from_policy(policy_text: str) -> list[str]:
    """Parse a style policy text and return detected required section/heading names."""
    return _extract_required_headings_from_policy_impl(policy_text)


def check_adf_against_style_policy(adf: dict, policy_text: str) -> list[str]:
    """Return warning strings for required headings that are absent in the ADF."""
    return _check_adf_against_style_policy(adf, policy_text, _node_text)


def run_set_policy(source: str, section: str | None = None):
    """Load a style policy from a local file or Confluence page URL/ID and save it."""
    print(f"\nLoading style policy from: {source}")
    if section:
        print(f"Extracting section: \"{section}\"")

    is_confluence_url = source.startswith('http')
    is_page_id = source.isdigit()

    if is_confluence_url or is_page_id:
        page_id = _page_id_from_url(source) if is_confluence_url else source
        if not page_id:
            print("ERROR: Could not extract page ID from URL.")
            print("  Expected: https://your-domain.atlassian.net/wiki/spaces/.../pages/123456/...")
            sys.exit(1)
        print(f"  Fetching Confluence page {page_id} ...")
        adf_body, page_title = fetch_page_adf(page_id)
        policy_text = _adf_to_markdown(adf_body)
        source_desc = page_title
    else:
        path = Path(source)
        if not path.exists():
            print(f"ERROR: File not found: {source}")
            sys.exit(1)
        ingested = ingest_file(source)
        policy_text = _adf_to_markdown(ingested) if isinstance(ingested, dict) else ingested
        source_desc = path.name

    if section:
        policy_text = _extract_section(policy_text, section)

    if not policy_text.strip():
        print("ERROR: No policy content found (empty after extraction).")
        sys.exit(1)

    preview = policy_text.strip().splitlines()
    print(f"\n── Preview ({'first 20 lines' if len(preview) > 20 else f'{len(preview)} lines'}) ──")
    for line in preview[:20]:
        print(f"  {line}")
    if len(preview) > 20:
        print(f"  ... ({len(preview) - 20} more lines not shown)")

    required = _extract_required_headings_from_policy(policy_text)
    if required:
        print(f"\n  Detected required sections ({len(required)}):")
        for section_name in required:
            print(f"    • {section_name}")
    else:
        print(
            "\n  ℹ  No structured required-section list detected.\n"
            "     Policy will be stored and Claude will apply it as contextual guidance."
        )

    print()
    confirm = input("Save as project style policy? [y/N] ").strip().lower()
    if confirm != 'y':
        print("Aborted.")
        return

    save_style_policy(policy_text, source=source_desc, section=section)
    print(
        f"\n✓ Policy active. Every document published or linted will be checked against it.\n"
        f"  View:  cat {_STYLE_POLICY_FILE}\n"
        f"  Clear: delete {_STYLE_POLICY_FILE}"
    )


def _normalize_tokens(text: str) -> set:
    """Lowercase, strip punctuation, remove common stop words — return set of tokens."""
    from publisherlib.policy import normalize_tokens as _normalize_tokens_impl

    return _normalize_tokens_impl(text)


def fuzzy_doc_match(title: str, regulation: str) -> tuple[str | None, float]:
    """Fuzzy-match a page title against a regulation's document catalog."""
    from publisherlib.policy import fuzzy_doc_match as _fuzzy_doc_match_impl

    return _fuzzy_doc_match_impl(title, regulation, REGULATION_CONFIGS)


def inject_regulation_doc_id(title: str, regulation: str | None) -> str:
    """Insert a matched regulation document ID into a title when applicable."""
    return _inject_regulation_doc_id(title, regulation, REGULATION_CONFIGS)


def strip_title_heading_from_adf(adf: dict, title: str) -> dict:
    """Remove the first ADF node if it is a heading that closely matches the page title."""
    return _strip_title_heading_from_adf(adf, title, _node_text)


def fix_adf_heading_numbers(adf: dict) -> tuple[dict, int]:
    """Strip numeric prefixes from all heading nodes in ADF."""
    return _fix_adf_heading_numbers(adf)


_DOC_LINT_CACHE = None


def find_doc_lint():
    """Search the Claude plugin cache for a doc-lint installation."""
    import glob as _glob

    patterns = [
        str(Path.home() / ".claude" / "plugins" / "cache" / "**" / "doc-lint" / "scripts" / "fix.py"),
        str(Path.home() / ".claude" / "plugins" / "**" / "doc-lint" / "scripts" / "fix.py"),
    ]
    for pattern in patterns:
        matches = _glob.glob(pattern, recursive=True)
        if matches:
            fix_py = Path(matches[0])
            lint_py = fix_py.parent / "lint.py"
            if lint_py.exists():
                return lint_py, fix_py
    return None, None


def get_doc_lint():
    """Return cached (lint_py, fix_py) or (None, None). Prints status on first call."""
    global _DOC_LINT_CACHE
    if _DOC_LINT_CACHE is None:
        lint_py, fix_py = find_doc_lint()
        _DOC_LINT_CACHE = (lint_py, fix_py)
        if lint_py:
            print(f"[doc-lint] Found — using enhanced cleanup rules from {fix_py.parent}")
        else:
            print("[doc-lint] Not found — using built-in cleanup rules (install doc-lint for enhanced rules)")
    return _DOC_LINT_CACHE


def get_headers():
    if not ATLASSIAN_EMAIL or not ATLASSIAN_TOKEN:
        print("ERROR: Missing ATLASSIAN_EMAIL or ATLASSIAN_API_TOKEN in .env")
        sys.exit(1)
    auth = base64.b64encode(f"{ATLASSIAN_EMAIL}:{ATLASSIAN_TOKEN}".encode()).decode()
    return {
        "Authorization": f"Basic {auth}",
        "Content-Type": "application/json",
        "Accept": "application/json",
    }


def list_spaces():
    return _list_spaces(ATLASSIAN_URL, get_headers())


def resolve_space_id(space_key: str) -> str:
    return _resolve_space_id(ATLASSIAN_URL, get_headers(), space_key)


def resolve_parent_id(space_key: str, parent_title: str) -> str | None:
    return _resolve_parent_id(ATLASSIAN_URL, get_headers(), space_key, parent_title)


def find_existing_page(space_id: str, title: str) -> dict | None:
    return _find_existing_page(ATLASSIAN_URL, get_headers(), space_id, title)


def get_page_version(page_id: str) -> int:
    return _get_page_version(ATLASSIAN_URL, get_headers(), page_id)


def list_child_pages(parent_id: str) -> list[dict]:
    return _list_child_pages(ATLASSIAN_URL, get_headers(), parent_id)


def delete_page(page_id: str):
    _delete_page(ATLASSIAN_URL, get_headers(), page_id)


def _google_doc_html_to_adf(html: str) -> dict:
    """Convert Google Docs HTML export to ADF, preserving headings, lists, tables, and inline marks."""
    from html.parser import HTMLParser as _HTMLParser

    class _Parser(_HTMLParser):
        def __init__(self):
            super().__init__(convert_charrefs=True)
            self.nodes = []
            self._buf = []
            self._marks = []
            self._mode = None
            self._hlevel = None
            self._lists = []
            self._table_rows = []
            self._row_buf = []
            self._cell_type = 'tableCell'

        def _style_marks(self, style):
            normalized = style.lower().replace(' ', '')
            marks = []
            if 'font-weight:bold' in normalized or 'font-weight:700' in normalized:
                marks.append({'type': 'strong'})
            if 'font-style:italic' in normalized:
                marks.append({'type': 'em'})
            if 'text-decoration:underline' in normalized:
                marks.append({'type': 'underline'})
            if 'text-decoration:line-through' in normalized:
                marks.append({'type': 'strike'})
            return marks

        def _flush(self):
            inline = []
            for text, marks in self._buf:
                if not text:
                    continue
                node = {'type': 'text', 'text': text}
                if marks:
                    node['marks'] = list(marks)
                inline.append(node)
            self._buf = []
            return inline

        def handle_starttag(self, tag, attrs):
            attr_dict = dict(attrs)
            if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
                self._mode = 'heading'
                self._hlevel = int(tag[1])
                self._buf = []
            elif tag == 'p':
                if self._mode != 'cell':
                    self._mode = 'para'
                self._buf = []
            elif tag in ('ul', 'ol'):
                self._lists.append((tag, []))
            elif tag == 'li':
                self._mode = 'li'
                self._buf = []
            elif tag == 'table':
                self._table_rows = []
            elif tag == 'tr':
                self._row_buf = []
            elif tag in ('th', 'td'):
                self._mode = 'cell'
                self._cell_type = 'tableHeader' if tag == 'th' else 'tableCell'
                self._buf = []
            elif tag in ('b', 'strong'):
                self._marks.append({'type': 'strong'})
            elif tag in ('i', 'em'):
                self._marks.append({'type': 'em'})
            elif tag == 'u':
                self._marks.append({'type': 'underline'})
            elif tag in ('s', 'strike'):
                self._marks.append({'type': 'strike'})
            elif tag == 'span':
                for mark in self._style_marks(attr_dict.get('style', '')):
                    self._marks.append(mark)

        def handle_endtag(self, tag):
            if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
                inline = self._flush()
                if inline:
                    self.nodes.append({'type': 'heading', 'attrs': {'level': self._hlevel}, 'content': inline})
                self._mode = None
            elif tag == 'p' and self._mode == 'para':
                inline = self._flush()
                if inline:
                    self.nodes.append({'type': 'paragraph', 'content': inline})
                self._mode = None
            elif tag == 'li':
                inline = self._flush()
                if inline and self._lists:
                    self._lists[-1][1].append({
                        'type': 'listItem',
                        'content': [{'type': 'paragraph', 'content': inline}],
                    })
                self._mode = None
            elif tag in ('ul', 'ol'):
                if self._lists:
                    list_type, items = self._lists.pop()
                    if items:
                        node_type = 'orderedList' if list_type == 'ol' else 'bulletList'
                        self.nodes.append({'type': node_type, 'content': items})
            elif tag in ('th', 'td'):
                inline = self._flush()
                self._row_buf.append({
                    'type': self._cell_type,
                    'attrs': {},
                    'content': [{'type': 'paragraph', 'content': inline}],
                })
                self._mode = None
            elif tag == 'tr':
                if self._row_buf:
                    self._table_rows.append({'type': 'tableRow', 'content': self._row_buf})
                self._row_buf = []
            elif tag == 'table':
                if self._table_rows:
                    self.nodes.append({
                        'type': 'table',
                        'attrs': {'isNumberColumnEnabled': False, 'layout': 'default'},
                        'content': self._table_rows,
                    })
            elif tag in ('b', 'strong', 'i', 'em', 'u', 's', 'strike', 'span'):
                target = {
                    'b': 'strong', 'strong': 'strong', 'i': 'em', 'em': 'em',
                    'u': 'underline', 's': 'strike', 'strike': 'strike',
                }.get(tag)
                for index in range(len(self._marks) - 1, -1, -1):
                    if target and self._marks[index].get('type') == target:
                        self._marks.pop(index)
                        break
                    if tag == 'span':
                        break

        def handle_data(self, data):
            if self._mode in ('heading', 'para', 'li', 'cell') and data:
                self._buf.append((data, list(self._marks)))

    parser = _Parser()
    parser.feed(html)
    return {'version': 1, 'type': 'doc', 'content': [node for node in parser.nodes if node]}


def ingest_google_doc(url: str) -> dict:
    """Export Google Doc as HTML and convert to ADF, preserving headings, lists, tables, and inline marks."""
    match = re.search(r"/document/d/([a-zA-Z0-9_-]+)", url)
    if not match:
        raise ValueError(f"Cannot extract doc ID from URL: {url}")
    doc_id = match.group(1)
    export_url = (
        f"https://docs.google.com/feeds/download/documents/export/Export"
        f"?id={doc_id}&exportFormat=html"
    )
    r = requests.get(export_url)
    if r.status_code == 403:
        raise PermissionError(
            f"Google Doc is not publicly accessible.\n"
            f"Share it as 'Anyone with the link can view' and retry.\n"
            f"URL: {url}"
        )
    r.raise_for_status()
    return _google_doc_html_to_adf(r.text)


def docx_to_adf(path: str) -> dict:
    """Convert a .docx file directly to ADF, preserving headings, lists, tables, and inline formatting."""
    try:
        from docx import Document
        from docx.oxml.ns import qn as _qn
        from docx.text.paragraph import Paragraph
        from docx.table import Table as DTable
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install -r requirements.txt")

    doc = Document(path)

    # ── numbering map ──────────────────────────────────────────────────────────
    def _num_type_map():
        try:
            numbering_el = doc.part.numbering_part._element
        except Exception:
            return {}
        abstract_nums = {}
        for an in numbering_el.findall(_qn('w:abstractNum')):
            an_id = an.get(_qn('w:abstractNumId'))
            for lvl in an.findall(_qn('w:lvl')):
                if lvl.get(_qn('w:ilvl')) == '0':
                    el = lvl.find(_qn('w:numFmt'))
                    if el is not None:
                        fmt = el.get(_qn('w:val'), '')
                        abstract_nums[an_id] = 'ordered' if fmt in (
                            'decimal', 'lowerLetter', 'upperLetter',
                            'lowerRoman', 'upperRoman') else 'bullet'
                    break
        num_map = {}
        for num in numbering_el.findall(_qn('w:num')):
            nid = num.get(_qn('w:numId'))
            ref = num.find(_qn('w:abstractNumId'))
            if ref is not None:
                num_map[nid] = abstract_nums.get(ref.get(_qn('w:val')), 'bullet')
        return num_map

    num_map = _num_type_map()

    def _get_numpr(pPr):
        if pPr is None:
            return None, None
        numPr = pPr.find(_qn('w:numPr'))
        if numPr is None:
            return None, None
        ilvl_el  = numPr.find(_qn('w:ilvl'))
        numid_el = numPr.find(_qn('w:numId'))
        ilvl  = ilvl_el.get(_qn('w:val'))  if ilvl_el  is not None else '0'
        numid = numid_el.get(_qn('w:val')) if numid_el is not None else None
        return numid, int(ilvl)

    def _para_size(para):
        for run in para.runs:
            if run.font.size:
                return run.font.size.pt
        try:
            sz = para.style.font.size
            if sz:
                return sz.pt
        except Exception:
            pass
        return None

    def _heading_level(para):
        style = para.style.name
        size  = _para_size(para)
        if style == 'Title':
            return 1
        if 'Heading 1' in style:
            return 2 if (size is None or size >= 13) else None
        if 'Heading 2' in style:
            return 3
        if 'Heading 3' in style:
            return 4
        if 'Heading 4' in style:
            return 5
        if style in ('Normal', 'Normal (Web)', 'Default Paragraph Style'):
            if size and size >= 18:
                return 1
            if size and size >= 13:
                return 2
        return None

    def _build_inline_with_breaks(para_el):
        """Walk paragraph XML, splitting into segments at <w:br> line-break elements."""
        current = []
        segments = [current]
        for child in para_el:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'r':
                rPr = child.find(_qn('w:rPr'))
                marks = []
                if rPr is not None:
                    if rPr.find(_qn('w:b'))      is not None: marks.append({'type': 'strong'})
                    if rPr.find(_qn('w:i'))      is not None: marks.append({'type': 'em'})
                    if rPr.find(_qn('w:u'))      is not None: marks.append({'type': 'underline'})
                    if rPr.find(_qn('w:strike')) is not None: marks.append({'type': 'strike'})
                for rchild in child:
                    rtag = rchild.tag.split('}')[-1] if '}' in rchild.tag else rchild.tag
                    if rtag == 't':
                        text = rchild.text or ''
                        if text:
                            node = {'type': 'text', 'text': text}
                            if marks: node['marks'] = marks[:]
                            current.append(node)
                    elif rtag == 'br':
                        current = []
                        segments.append(current)
        return [s for s in segments if s]

    def _build_inline(para):
        nodes = []
        for run in para.runs:
            if not run.text:
                continue
            marks = []
            if run.bold:      marks.append({'type': 'strong'})
            if run.italic:    marks.append({'type': 'em'})
            if run.underline: marks.append({'type': 'underline'})
            node = {'type': 'text', 'text': run.text}
            if marks:
                node['marks'] = marks
            nodes.append(node)
        if not nodes and para.text.strip():
            nodes.append({'type': 'text', 'text': para.text})
        return nodes

    def _consecutive_headings_to_lists(nodes):
        """Convert runs of 3+ consecutive same-level headings (level >= 3) to bullet lists."""
        result = []
        i = 0
        while i < len(nodes):
            node = nodes[i]
            if node['type'] != 'heading' or node['attrs']['level'] < 3:
                result.append(node)
                i += 1
                continue
            level = node['attrs']['level']
            run = [node]
            j = i + 1
            while j < len(nodes) and nodes[j]['type'] == 'heading' and nodes[j]['attrs']['level'] == level:
                run.append(nodes[j])
                j += 1
            if len(run) >= 3:
                items = [{'type': 'listItem', 'content': [{'type': 'paragraph', 'content': h['content']}]}
                         for h in run]
                result.append({'type': 'bulletList', 'content': items})
                i = j
            else:
                result.append(node)
                i += 1
        return result

    def _convert_table(tbl_el):
        table = DTable(tbl_el, doc)
        rows = []
        for ri, row in enumerate(table.rows):
            cells = []
            seen = set()
            for cell in row.cells:
                cid = id(cell)
                if cid in seen:
                    continue
                seen.add(cid)
                cell_content = []
                for p in cell.paragraphs:
                    inline = _build_inline(p)
                    if inline:
                        cell_content.append({'type': 'paragraph', 'content': inline})
                if not cell_content:
                    cell_content = [{'type': 'paragraph', 'content': [{'type': 'text', 'text': ''}]}]
                ctype = 'tableHeader' if ri == 0 else 'tableCell'
                cells.append({'type': ctype, 'attrs': {}, 'content': cell_content})
            if cells:
                rows.append({'type': 'tableRow', 'content': cells})
        return {'type': 'table', 'attrs': {'isNumberColumnEnabled': False, 'layout': 'default'}, 'content': rows}

    # ── main pass ──────────────────────────────────────────────────────────────
    nodes = []
    pending_list = None  # {'type': 'bullet'|'ordered', 'items': [...]}

    def _flush():
        nonlocal pending_list
        if pending_list:
            lnode = 'orderedList' if pending_list['type'] == 'ordered' else 'bulletList'
            nodes.append({'type': lnode, 'content': pending_list['items']})
            pending_list = None

    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

        if tag == 'tbl':
            _flush()
            nodes.append(_convert_table(element))

        elif tag == 'p':
            para   = Paragraph(element, doc)
            pPr    = element.find(_qn('w:pPr'))
            numid, ilvl = _get_numpr(pPr)

            if numid and numid != '0':
                ltype  = num_map.get(numid, 'bullet')
                inline = _build_inline(para)
                if not inline:
                    continue
                item = {'type': 'listItem', 'content': [{'type': 'paragraph', 'content': inline}]}
                if pending_list and pending_list['type'] == ltype:
                    pending_list['items'].append(item)
                else:
                    _flush()
                    pending_list = {'type': ltype, 'items': [item]}
            else:
                _flush()
                if not para.text.strip():
                    continue
                lvl      = _heading_level(para)
                segments = _build_inline_with_breaks(element)
                if lvl and len(segments) > 1:
                    # Multiline heading paragraph — split at line breaks
                    nodes.append({'type': 'heading', 'attrs': {'level': lvl}, 'content': segments[0]})
                    for seg in segments[1:]:
                        nodes.append({'type': 'paragraph', 'content': seg})
                elif lvl:
                    inline = segments[0] if segments else _build_inline(para)
                    nodes.append({'type': 'heading', 'attrs': {'level': lvl}, 'content': inline})
                else:
                    # Flatten segments (soft returns in body text become spaces)
                    inline = []
                    for si, seg in enumerate(segments):
                        inline.extend(seg)
                        if si < len(segments) - 1:
                            inline.append({'type': 'text', 'text': ' '})
                    if not inline:
                        inline = [{'type': 'text', 'text': para.text}]
                    nodes.append({'type': 'paragraph', 'content': inline})

    _flush()
    nodes = _consecutive_headings_to_lists(nodes)
    return {'version': 1, 'type': 'doc', 'content': nodes}


def ingest_pdf(path: str) -> str:
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber not installed. Run: pip install -r requirements.txt")
    with pdfplumber.open(path) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages)


def ingest_file(source: str):
    """Return ADF dict for .docx/.json ADF, or plain text str for all other sources."""
    if source.startswith("https://docs.google.com"):
        return ingest_google_doc(source)
    p = Path(source)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {source}")
    ext = p.suffix.lower()
    if ext == ".docx":
        _, fix_py = get_doc_lint()
        if fix_py:
            # Run doc-lint fix.py on a temp copy, then convert the cleaned file
            import subprocess, shutil
            tmp = p.with_suffix("._tmp_.docx")
            shutil.copy(str(p), str(tmp))
            try:
                result = subprocess.run(
                    [sys.executable, str(fix_py), "--file", str(tmp), "--overwrite"],
                    capture_output=True,
                    text=True,
                )
                if result.returncode != 0:
                    _warn(f"doc-lint cleanup failed for {p.name}; falling back to built-in cleanup")
                    if result.stderr:
                        print(result.stderr.strip(), file=sys.stderr)
                    return docx_to_adf(str(p))
                return docx_to_adf(str(tmp))
            except Exception as exc:
                _warn(f"doc-lint cleanup crashed for {p.name}; falling back to built-in cleanup: {exc}")
                return docx_to_adf(str(p))
            finally:
                tmp.unlink(missing_ok=True)
        return docx_to_adf(str(p))   # built-in cleanup via docx_to_adf
    elif ext == ".pdf":
        return ingest_pdf(str(p))
    elif ext == ".json":
        try:
            data = json.loads(p.read_text(encoding='utf-8'))
            # Accept bare ADF doc or wrapped {"adf": {...}} / {"body": {...}} shapes
            if isinstance(data, dict):
                if data.get('type') == 'doc' and 'content' in data:
                    return data
                for key in ('adf', 'body', 'value'):
                    if isinstance(data.get(key), dict) and data[key].get('type') == 'doc':
                        return data[key]
            raise ValueError("JSON file does not look like an ADF document")
        except (json.JSONDecodeError, ValueError) as e:
            raise ValueError(f"Cannot read {p.name} as ADF JSON: {e}") from e
    else:
        return p.read_text(encoding="utf-8")

def _adf_to_markdown(adf: dict) -> str:
    """Convert an ADF document to a markdown-like plain text string.

    Preserves heading levels as # markers so that _extract_section() can
    find and extract named sections from ADF-sourced documents.
    """
    lines: list[str] = []

    def walk(node: dict, list_marker: str = ''):
        ntype = node.get('type', '')
        children = node.get('content', [])

        if ntype == 'heading':
            level = node.get('attrs', {}).get('level', 1)
            text  = _node_text(node).strip()
            lines.append('#' * level + ' ' + text)
        elif ntype in ('paragraph', 'tableCell', 'tableHeader'):
            text = _node_text(node).strip()
            if text:
                lines.append(list_marker + text if list_marker else text)
        elif ntype == 'listItem':
            # walk items with a bullet marker, passing it down to paragraphs
            for child in children:
                walk(child, list_marker='- ')
            return
        elif ntype in ('bulletList', 'orderedList'):
            for child in children:
                walk(child)
            return
        elif ntype == 'rule':
            lines.append('---')
        elif ntype == 'panel':
            text = _node_text(node).strip()
            if text:
                lines.append(f'> {text}')
        elif ntype in ('table', 'tableRow'):
            for child in children:
                walk(child)
            return
        elif ntype in ('doc', 'blockquote', 'expand'):
            for child in children:
                walk(child)
            return
        else:
            # Unknown block — walk children
            for child in children:
                walk(child)

    for node in adf.get('content', []):
        walk(node)

    return '\n'.join(lines)


# ---------------------------------------------------------------------------
# ADF builder
# ---------------------------------------------------------------------------

def text_node(text: str, marks: list = None) -> dict:
    node = {"type": "text", "text": text}
    if marks:
        node["marks"] = marks
    return node


def paragraph(*texts) -> dict:
    return {"type": "paragraph", "content": [text_node(t) for t in texts]}


def heading(level: int, text: str) -> dict:
    return {
        "type": "heading",
        "attrs": {"level": level},
        "content": [text_node(text)],
    }


def bullet_list(items: list[str]) -> dict:
    return {
        "type": "bulletList",
        "content": [
            {"type": "listItem", "content": [paragraph(item)]}
            for item in items
        ],
    }


def info_panel(text: str, panel_type: str = "info") -> dict:
    return {
        "type": "panel",
        "attrs": {"panelType": panel_type},
        "content": [paragraph(text)],
    }


def metadata_table(rows: list[tuple[str, str]]) -> dict:
    def header_cell(text):
        return {"type": "tableHeader", "content": [paragraph(text)]}
    def data_cell(text):
        return {"type": "tableCell", "content": [paragraph(text)]}

    return {
        "type": "table",
        "attrs": {"isNumberColumnEnabled": False, "layout": "default"},
        "content": [
            {
                "type": "tableRow",
                "content": [header_cell(k), data_cell(v)],
            }
            for k, v in rows
        ],
    }


def to_adf(content: str, template: str = "general") -> dict:
    """
    Convert plain text content to ADF, applying template structure.
    This is a scaffold — Claude Code will expand section detection per template.
    """
    lines = content.strip().splitlines()
    nodes = []

    # Template metadata headers
    if template == "iso27001":
        nodes.append(info_panel(
            "ISO 27001 Compliance Document \u2014 review Annex A control mapping before approving.",
            "warning"
        ))
        # Note: document title, version, and approval metadata are managed
        # by the eSign Document plugin. Only classification is added inline.

    # Convert source lines to ADF nodes
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # Wiki markup heading detection
        if line.startswith("h1. "):
            nodes.append(heading(1, line[4:]))
        elif line.startswith("h2. "):
            nodes.append(heading(2, line[4:]))
        elif line.startswith("h3. "):
            nodes.append(heading(3, line[4:]))
        elif line.startswith("# ") or line.startswith("## ") or line.startswith("### "):
            level = len(line.split(" ")[0])
            nodes.append(heading(level, line.lstrip("# ")))
        elif line.startswith("* ") or line.startswith("- "):
            nodes.append(bullet_list([line[2:]]))
        else:
            nodes.append(paragraph(line))

    # Template section scaffolding — add missing required sections
    section_map = {
        "policy": [
            "Purpose", "Scope", "Definitions", "Roles and Responsibilities",
            "Policy Statements", "Compliance and Exceptions", "Related Documents", "Revision History"
        ],
        "procedure": [
            "Purpose", "Scope", "Prerequisites", "Procedure Steps",
            "Exceptions and Escalations", "Related Documents", "Revision History"
        ],
        "workflow": [
            "Purpose", "Trigger", "Roles Involved", "Flow Steps",
            "Decision Points", "Outcomes", "Related Documents"
        ],
        "record": ["Metadata", "Data Fields", "Audit Trail"],
        "form": ["Instructions", "Fields", "Submission Guidance"],
        "meeting_minutes": ["Attendees", "Agenda", "Discussion", "Decisions", "Action Items"],
        "iso27001": [
            "Purpose", "Scope", "Definitions", "Roles and Responsibilities",
            "Policy Statements", "Control Mapping", "Compliance and Exceptions",
            "Related Documents", "Revision History"
        ],
    }

    required = section_map.get(template.lower().replace(" ", "_"), [])
    existing_text = content.lower()
    for section in required:
        if section.lower() not in existing_text:
            nodes.append(heading(2, section))
            nodes.append(paragraph(f"[TO BE COMPLETED — {section}]"))

    return {"version": 1, "type": "doc", "content": nodes}

# ---------------------------------------------------------------------------
# Publish
# ---------------------------------------------------------------------------

def create_page(
    space_id: str,
    title: str,
    adf_body: dict,
    parent_id: str = None,
    labels: list[str] = None,
    status: str = "current",
) -> dict:
    return _create_page(
        ATLASSIAN_URL,
        get_headers(),
        space_id,
        title,
        adf_body,
        parent_id=parent_id,
        labels=labels,
        status=status,
    )


def update_page(
    page_id: str,
    title: str,
    adf_body: dict,
    status: str = "current",
) -> dict:
    version = get_page_version(page_id)
    return _update_page(
        ATLASSIAN_URL,
        get_headers(),
        page_id,
        title,
        adf_body,
        version,
        status=status,
    )


def page_url(page: dict) -> str:
    return _page_url(ATLASSIAN_URL, page)

# ---------------------------------------------------------------------------
# Mapping file loader
# ---------------------------------------------------------------------------

def load_mapping(path: str) -> list[dict]:
    try:
        import pandas as pd
    except ImportError:
        raise ImportError("pandas not installed. Run: pip install -r requirements.txt")

    ext = Path(path).suffix.lower()
    if ext in (".xlsx", ".xls"):
        df = pd.read_excel(path, dtype=str)
    else:
        df = pd.read_csv(path, dtype=str)

    df = df.fillna("")
    required = {"file", "space_key"}
    missing = required - set(df.columns.str.lower())
    if missing:
        raise ValueError(f"Mapping file missing required columns: {missing}")

    df.columns = df.columns.str.lower().str.strip()
    return df.to_dict(orient="records")

# ---------------------------------------------------------------------------
# Confluence audit and remediation
# ---------------------------------------------------------------------------

def build_space_tree(space_key: str) -> list:
    """Build a full hierarchical tree of pages and folders in a space.

    Uses CQL with expand=ancestors so every item knows its direct parent.
    Returns a list of root-level node dicts; each node has:
      {id, title, type, children: [...]}
    """
    return _build_space_tree(ATLASSIAN_URL, get_headers(), space_key)


def _print_tree_nodes(nodes: list, indent: int = 0):
    """Recursively print a tree produced by build_space_tree()."""
    for node in sorted(nodes, key=lambda n: n["title"]):
        icon   = "📁" if node["type"] == "folder" else "📄"
        prefix = "  " * indent
        print(f"{prefix}{icon} {node['title']}")
        if node["children"]:
            _print_tree_nodes(node["children"], indent + 1)


def _count_tree(nodes: list) -> tuple[int, int]:
    """Return (page_count, folder_count) for a tree."""
    pages, folders = 0, 0
    for n in nodes:
        if n["type"] == "folder":
            folders += 1
        else:
            pages += 1
        cp, cf = _count_tree(n["children"])
        pages += cp
        folders += cf
    return pages, folders


def fetch_page_adf(page_id: str) -> tuple:
    """Fetch a Confluence page and return (adf_body, title)."""
    return _fetch_page_adf(ATLASSIAN_URL, get_headers(), page_id)


def walk_descendant_pages(parent_id: str) -> list:
    """Return all descendant pages under parent_id using CQL ancestor search."""
    return _walk_descendant_pages(ATLASSIAN_URL, get_headers(), parent_id)


def walk_space_pages(space_key: str) -> list:
    """Return all pages in a space using CQL."""
    return _walk_space_pages(ATLASSIAN_URL, get_headers(), space_key)


def remediate_adf(adf: dict, template: str, missing_sections: list) -> dict:
    """
    Insert a warning panel + placeholder paragraph for each missing required section.
    Inserts before 'Revision History' if present; otherwise appends at end.
    Returns a new ADF dict (does not mutate the input).
    """
    import copy
    new_adf = copy.deepcopy(adf)
    nodes   = new_adf.setdefault("content", [])

    # Find insertion point — before 'Revision History' if present
    insert_idx = len(nodes)
    for i, node in enumerate(nodes):
        if node.get("type") == "heading":
            text = "".join(
                c.get("text", "") for c in node.get("content", [])
                if c.get("type") == "text"
            )
            if "revision history" in text.lower():
                insert_idx = i
                break

    # Order missing sections by their position in the template's required order
    required_order = REQUIRED_SECTIONS.get(template.lower(), [])
    ordered = [s for s in required_order if s in missing_sections]
    ordered += [s for s in missing_sections if s not in ordered]  # catch any extras

    new_nodes = []
    for section in ordered:
        new_nodes.append(heading(2, section))
        new_nodes.append({
            "type": "panel",
            "attrs": {"panelType": "warning"},
            "content": [paragraph(f"[TO BE COMPLETED — {section}]")],
        })

    nodes[insert_idx:insert_idx] = new_nodes
    return new_adf


def run_audit(space_key: str, folder: str = None) -> list:
    """
    Scan all pages in a space (or under a named folder/page) and audit each for
    template compliance, missing required sections, and naming convention.
    Returns list of audit result dicts.
    """
    if folder:
        parent_id = resolve_parent_id(space_key, folder)
        print(f"\nScanning '{folder}' in space {space_key}...")
        pages = walk_descendant_pages(parent_id)
        # Include the root folder/parent page itself in the audit
        pages = [{"id": parent_id, "title": folder}] + pages
    else:
        print(f"\nScanning entire space {space_key}...")
        pages = walk_space_pages(space_key)

    if not pages:
        print("No pages found.")
        return []

    print(f"Found {len(pages)} page(s). Auditing...\n")

    results = []
    for i, page in enumerate(pages, 1):
        page_id    = page["id"]
        page_title = page.get("title", page_id)
        # Overwrite line with progress indicator
        print(f"  [{i:>3}/{len(pages)}] {page_title[:65]:<65}", end="\r", flush=True)

        try:
            adf, _   = fetch_page_adf(page_id)
            text     = _extract_text_from_adf(adf.get("content", []))
            template = detect_template_from_text(text)
            missing  = check_template_sections(adf["content"], template)
            name_ok, name_ex = validate_naming_convention(page_title, template)
            results.append({
                "id":       page_id,
                "title":    page_title,
                "template": template,
                "missing":  missing,
                "name_ok":  name_ok,
                "name_ex":  name_ex,
                "adf":      adf,
                "compliant": not missing and name_ok,
            })
        except Exception as e:
            results.append({
                "id": page_id, "title": page_title, "template": "?",
                "missing": [], "name_ok": True, "name_ex": "",
                "adf": None, "compliant": None, "error": str(e),
            })

    print()  # clear the carriage-return progress line
    return results


def print_audit_report(results: list, space_key: str):
    """Print a human-readable compliance audit report."""
    compliant     = [r for r in results if r.get("compliant") is True]
    non_compliant = [r for r in results if r.get("compliant") is False]
    errors        = [r for r in results if "error" in r]

    print(f"\n{'='*70}")
    print(f"Audit report — {space_key}")
    print(f"{'='*70}")
    print(f"  Total pages   : {len(results)}")
    print(f"  ✓  Compliant  : {len(compliant)}")
    print(f"  ✗  Issues     : {len(non_compliant)}")
    if errors:
        print(f"  ⚠  Errors     : {len(errors)}")

    if non_compliant:
        print(f"\n{'─'*70}")
        print("Non-compliant pages:")
        print(f"{'─'*70}")
        for r in non_compliant:
            print(f"\n  {r['title']}")
            print(f"    Template : {r['template']}")
            if r["missing"]:
                print(f"    Missing  : {', '.join(r['missing'])}")
            if not r["name_ok"] and r["name_ex"]:
                print(f"    Naming   : expected — {r['name_ex']}")

    if errors:
        print(f"\n{'─'*70}")
        print("Errors (pages skipped):")
        for r in errors:
            print(f"  {r['title']}: {r.get('error', '')}")

    print()


def run_remediate(space_key: str, folder: str = None, go: bool = False):
    """
    Audit all pages in the space/folder, then patch each non-compliant page
    by inserting warning-panel placeholders for every missing required section.
    Naming convention violations are reported but not auto-fixed (require a rename).
    """
    results = run_audit(space_key, folder)
    if not results:
        return

    print_audit_report(results, space_key)

    # Pages with missing sections — these can be auto-remediated
    fixable = [
        r for r in results
        if r.get("compliant") is False
        and r["adf"] is not None
        and r["missing"]
    ]
    # Pages with naming violations only — report at end
    naming_only = [
        r for r in results
        if r.get("compliant") is False
        and not r["missing"]
        and not r["name_ok"]
    ]

    if not fixable:
        print("No missing sections detected — nothing to auto-remediate.")
    else:
        print(f"Remediation plan — {len(fixable)} page(s):")
        print(f"{'─'*70}")
        for r in fixable:
            print(f"  {r['title']}")
            print(f"    Add: {', '.join(r['missing'])}")
        print()

        if not go:
            confirm = input(f"Proceed? [y/N] ").strip().lower()
            if confirm != "y":
                print("Aborted.")
                return

        ok = fail = 0
        for r in fixable:
            try:
                new_adf = remediate_adf(r["adf"], r["template"], r["missing"])
                update_page(r["id"], r["title"], new_adf)
                print(f"  ✓  {r['title']}")
                ok += 1
            except Exception as e:
                print(f"  ✗  {r['title']} — {e}")
                fail += 1

        print(f"\nRemediation complete: {ok} updated, {fail} failed")

    # Always report naming violations — these need manual renames
    naming_issues = [
        r for r in results
        if not r.get("name_ok") and r.get("template") not in ("general", "?") and r.get("name_ex")
    ]
    if naming_issues:
        print(f"\nNaming violations (manual rename required — {len(naming_issues)} page(s)):")
        print(f"{'─'*70}")
        for r in naming_issues:
            print(f"  '{r['title']}'  →  expected: {r['name_ex']}")
        print()


def run_fix_heading_numbers(space_key: str, folder: str = None, go: bool = False):
    """Scan all pages in a space (or folder) for numbered heading prefixes and strip them."""
    print(f"\nScanning for numbered heading prefixes in {space_key}" +
          (f" / '{folder}'" if folder else "") + " ...")

    if folder:
        parent_id = resolve_parent_id(space_key, folder)
        pages = walk_descendant_pages(parent_id)
    else:
        pages = walk_space_pages(space_key)

    plan = []  # list of {page_id, title, count, new_adf}
    for page in pages:
        adf, title = fetch_page_adf(page['id'])
        fixed_adf, count = fix_adf_heading_numbers(adf)
        if count > 0:
            plan.append({'page_id': page['id'], 'title': title,
                         'count': count, 'new_adf': fixed_adf})

    if not plan:
        print("\n\u2713 No numbered heading prefixes found.")
        return

    print(f"\nFix plan \u2014 {len(plan)} page(s) with numbered heading prefixes:\n")
    print("\u2500" * 60)
    for p in plan:
        print(f"  {p['title']}")
        print(f"    {p['count']} heading(s) will have number prefix stripped")
    print()

    if not go:
        confirm = input("Proceed? [y/N] ").strip().lower()
        if confirm != 'y':
            print("Aborted.")
            return

    ok, fail = 0, 0
    for p in plan:
        try:
            update_page(p['page_id'], p['title'], p['new_adf'])
            print(f"  \u2713  {p['title']}")
            ok += 1
        except Exception as e:
            print(f"  \u2717  {p['title']}: {e}")
            fail += 1

    print(f"\nDone: {ok} fixed, {fail} failed")


def run_add_print_headers(space_key: str, folder: str = None, go: bool = False):
    """Add document control header and print footer to pages that are missing them."""
    print(f"\nScanning for pages without document control blocks in {space_key}" +
          (f" / '{folder}'" if folder else "") + " ...")

    if folder:
        parent_id = resolve_parent_id(space_key, folder)
        pages = walk_descendant_pages(parent_id)
    else:
        pages = walk_space_pages(space_key)

    plan = []
    for page in pages:
        adf, title = fetch_page_adf(page['id'])
        if _has_doc_control_header(adf):
            continue
        doc_id = extract_doc_id_from_title(title)
        new_adf = wrap_with_print_controls(adf, title, doc_id=doc_id)
        plan.append({'page_id': page['id'], 'title': title, 'new_adf': new_adf})

    if not plan:
        print("\n\u2713 All pages already have document control blocks.")
        return

    print(f"\nFix plan \u2014 {len(plan)} page(s) missing document control header/footer:\n")
    print("\u2500" * 60)
    for p in plan:
        print(f"  {p['title']}")
    print()

    if not go:
        confirm = input("Proceed? [y/N] ").strip().lower()
        if confirm != 'y':
            print("Aborted.")
            return

    ok, fail = 0, 0
    for p in plan:
        try:
            update_page(p['page_id'], p['title'], p['new_adf'])
            print(f"  \u2713  {p['title']}")
            ok += 1
        except Exception as e:
            print(f"  \u2717  {p['title']}: {e}")
            fail += 1

    print(f"\nDone: {ok} updated, {fail} failed")


def _prepare_adf_for_publish(source: str, template: str, title: str) -> tuple[dict, str]:
    content = ingest_file(source)
    adf = content if isinstance(content, dict) else to_adf(content, template)
    json.loads(json.dumps(adf))

    adf = strip_title_heading_from_adf(adf, title)
    adf, heading_fixes = fix_adf_heading_numbers(adf)
    if heading_fixes:
        print(f"  Stripped number prefix from {heading_fixes} heading(s)")

    adf = wrap_with_print_controls(adf, title, doc_id=extract_doc_id_from_title(title))

    policy_text, _ = load_style_policy()
    if policy_text:
        for warning in check_adf_against_style_policy(adf, policy_text):
            print(f"  ⚠  {warning}")

    source_text = _extract_text_from_adf(adf.get("content", []))
    detected_template = detect_template_from_text(source_text) if template == "general" else template
    return adf, detected_template


def _apply_regulation_title(title: str) -> str:
    reg_cfg = load_regulation_config()
    regulation = reg_cfg.get("regulation")
    if not regulation:
        return title
    new_title = inject_regulation_doc_id(title, regulation)
    if new_title != title:
        print(f"[{REGULATION_CONFIGS[regulation]['name']}] Title updated: {title} → {new_title}")
    return new_title


def _print_compliance_warnings(source: str, detected_template: str, adf: dict):
    missing_sections = check_template_sections(adf["content"], detected_template)
    name_valid, name_example = validate_naming_convention(Path(source).stem, detected_template)
    if not name_valid and not missing_sections:
        print(f"\nCompliance warnings (template: {detected_template}):")
    elif not name_valid or missing_sections:
        print(f"\nCompliance warnings (template: {detected_template}):")
    if not name_valid:
        print(f"  ⚠  Naming convention: '{Path(source).stem}' does not match expected pattern")
        print(f"     Expected format: {name_example}")
    for section in missing_sections:
        print(f"  ⚠  Missing required section: '{section}'")


def _publish_single_file(args):
    if not args.space:
        print("ERROR: --space required for single file mode")
        sys.exit(1)

    print(f"Ingesting: {args.file}")
    title = _apply_regulation_title(args.title or normalize_title(Path(args.file).stem))
    adf, detected_template = _prepare_adf_for_publish(args.file, args.template, title)
    _print_compliance_warnings(args.file, detected_template, adf)

    labels = [label.strip() for label in args.labels.split(",")] if args.labels else []

    if args.preview:
        import tempfile
        import webbrowser

        html = adf_to_html(adf, title)
        tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".html", delete=False, encoding="utf-8")
        tmp.write(html)
        tmp.close()
        webbrowser.open(f"file://{tmp.name}")
        print(f"\n[Preview] Opened in browser: {tmp.name}")
        go = input("Continue to publish? [y/N] ").strip().lower()
        if go != "y":
            print(f"Aborted. Preview saved at: {tmp.name}")
            return

    print(f"\nUpload plan:")
    print(f"  Title:    {title}")
    print(f"  Space:    {args.space}")
    print(f"  Parent:   {args.parent or '(space root)'}")
    print(f"  Template: {args.template}")
    print(f"  Labels:   {', '.join(labels) or 'none'}")

    if args.dry_run:
        print("\n[DRY RUN] Would publish the above. Done.")
        return

    if not args.go:
        confirm = input("\nProceed? [y/N] ").strip().lower()
        if confirm != "y":
            print("Aborted.")
            return

    space_id = resolve_space_id(args.space)
    parent_id = resolve_parent_id(args.space, args.parent) if args.parent else None
    existing = find_existing_page(space_id, title)

    if existing:
        live_adf, _ = fetch_page_adf(existing["id"])
        changed, diff_summary = diff_adf(live_adf, adf)
        if not changed:
            print(f"\nNo changes detected — '{title}' is already up to date. Skipped.")
            return
        print(f"\nUpdating '{title}':")
        for line in diff_summary:
            print(line)
        page = update_page(existing["id"], title, adf)
    else:
        page = create_page(space_id, title, adf, parent_id, labels)

    print(f"\nPublished: {page_url(page)}")


def _publish_mapping(args):
    rows = load_mapping(args.mapping)
    print(f"\nLoaded {len(rows)} rows from mapping file.")

    print(f"\n{'FILE':<40} {'TITLE':<40} {'SPACE':<8} {'PARENT':<30} {'TEMPLATE':<15}")
    print("-" * 140)
    for row in rows:
        print(
            f"{row.get('file',''):<40} {row.get('title','(auto)'):<40} "
            f"{row.get('space_key',''):<8} {row.get('parent_page','(root)'):<30} "
            f"{row.get('template', args.template):<15}"
        )

    if args.dry_run:
        print("\n[DRY RUN] Would publish the above. Done.")
        return

    if not args.go:
        confirm = input(f"\nProceed with publishing {len(rows)} pages? [y/N] ").strip().lower()
        if confirm != "y":
            print("Aborted.")
            return

    collision_default = "1" if args.go else None
    apply_to_all = args.go
    results = []

    for row in rows:
        source = row.get("file", "")
        title = row.get("title", "").strip() or normalize_title(Path(source).stem)
        try:
            print(f"\nProcessing: {source}")
            title = _apply_regulation_title(title)
            template = row.get("template", args.template) or args.template
            adf, detected_template = _prepare_adf_for_publish(source, template, title)

            missing_sections = check_template_sections(adf["content"], detected_template)
            name_valid, name_example = validate_naming_convention(Path(source).stem, detected_template)
            if not name_valid:
                print(f"  ⚠  Naming: '{Path(source).stem}' doesn't match {detected_template} pattern ({name_example})")
            for section in missing_sections:
                print(f"  ⚠  Missing section: '{section}'")

            space_key = row.get("space_key", args.space or "")
            parent_title = row.get("parent_page", "")
            raw_labels = row.get("labels", "")
            labels = [label.strip() for label in raw_labels.split(",")] if raw_labels else []
            space_id = resolve_space_id(space_key)
            parent_id = resolve_parent_id(space_key, parent_title) if parent_title else None
            existing = find_existing_page(space_id, title)

            if existing:
                if apply_to_all and collision_default:
                    choice = collision_default
                else:
                    print(f"  COLLISION: '{title}' exists.")
                    raw = input("  1=Overwrite  2=Skip  3=Cancel all  Apply to all? [y/n]: ")
                    parts = raw.strip().split()
                    choice = parts[0] if parts else "2"
                    if len(parts) > 1 and parts[1].lower() == "y":
                        apply_to_all = True
                        collision_default = choice

                if choice == "3":
                    print("Cancelled.")
                    break
                if choice == "2":
                    results.append({"file": source, "title": title, "status": "SKIP", "url": ""})
                    continue

                live_adf, _ = fetch_page_adf(existing["id"])
                changed, diff_summary = diff_adf(live_adf, adf)
                if not changed:
                    results.append({"file": source, "title": title, "status": "UNCHANGED", "url": page_url(existing)})
                    print("  Unchanged — skipped.")
                    continue
                for line in diff_summary:
                    print(f" {line}")
                page = update_page(existing["id"], title, adf)
            else:
                page = create_page(space_id, title, adf, parent_id, labels)

            url = page_url(page)
            results.append({"file": source, "title": title, "status": "OK", "url": url})
            print(f"  Published: {url}")
        except Exception as exc:
            results.append({"file": source, "title": title, "status": "FAIL", "url": str(exc)})
            print(f"  FAILED: {exc}")

    ok = sum(1 for result in results if result["status"] == "OK")
    skip = sum(1 for result in results if result["status"] == "SKIP")
    unchanged = sum(1 for result in results if result["status"] == "UNCHANGED")
    fail = sum(1 for result in results if result["status"] == "FAIL")

    print(f"\n{'='*60}")
    print(f"Results: {ok} published, {unchanged} unchanged, {skip} skipped, {fail} failed")
    print(f"{'FILE':<40} {'STATUS':<8} {'URL'}")
    print("-" * 100)
    for result in results:
        print(f"{result['file']:<40} {result['status']:<8} {result['url']}")


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

def analyze_file(source: str):
    """Run structural analysis on a file and print a report. No publishing."""
    p = Path(source)

    # Delegate structural checks to doc-lint when available (richer output)
    lint_py, _ = get_doc_lint()
    if lint_py and p.suffix.lower() == '.docx':
        import subprocess
        result = subprocess.run(
            [sys.executable, str(lint_py), '--file', str(p)],
            capture_output=True, text=True,
        )
        if result.stdout:
            print(result.stdout)
        if result.stderr:
            print(result.stderr, file=sys.stderr)
    else:
        # Built-in fallback checks
        from docx import Document
        from docx.oxml.ns import qn

        def _para_size(para):
            for run in para.runs:
                if run.font.size: return run.font.size.pt
            return None

        def _heading_level(para):
            style = para.style.name
            size  = _para_size(para)
            if style == 'Title': return 1
            if 'Heading 1' in style: return 2 if (size is None or size >= 13) else None
            if 'Heading 2' in style: return 3
            if 'Heading 3' in style: return 4
            if style in ('Normal', 'Normal (Web)'):
                if size and size >= 18: return 1
                if size and size >= 13: return 2
            return None

        doc = Document(str(p))
        issues = []
        headings = paras = lists = 0
        consec = 0
        misuse = 0

        for para in doc.paragraphs:
            if not para.text.strip(): continue
            pPr = para._element.find(qn('w:pPr'))
            numPr = pPr.find(qn('w:numPr')) if pPr is not None else None
            if numPr is not None:
                lists += 1; consec = 0; continue
            lvl = _heading_level(para)
            style = para.style.name
            size  = _para_size(para)
            if 'Heading 1' in style and size and size < 13:
                misuse += 1
            if lvl:
                headings += 1; consec += 1
                if consec >= 3:
                    issues.append(f'⚠  Consecutive headings: "{para.text.strip()[:50]}" is heading #{consec} in a row')
            else:
                paras += 1; consec = 0

        if misuse:
            issues.append(f'⚠  Style misuse: {misuse} "Heading 1" paragraphs at ≤12pt — reclassify as body text')

        print(f"\nAnalysis: {source}")
        print("─" * 60)
        print(f"Structure: {headings} headings, {paras} paragraphs, {lists} list items, {len(doc.tables)} tables")
        if issues:
            print(f"\nIssues ({len(issues)}):")
            for issue in issues:
                print(f"  {issue}")
        else:
            print("\nIssues: none")

    # Template detection and compliance (always runs)
    if p.suffix.lower() == '.docx':
        from docx import Document
        doc = Document(str(p))
        source_text = ' '.join(para.text for para in doc.paragraphs)
        adf = docx_to_adf(str(p))
    else:
        source_text = p.read_text(encoding='utf-8')
        adf = to_adf(source_text)

    detected_template = detect_template_from_text(source_text)
    missing_sections  = check_template_sections(adf['content'], detected_template)
    name_valid, name_example = validate_naming_convention(p.stem, detected_template)

    print(f"\nTemplate (auto-detected): {detected_template}")
    if missing_sections:
        print(f"Missing required sections ({len(missing_sections)}): {', '.join(missing_sections)}")
    else:
        print("Required sections: all present ✓")
    if not name_valid:
        print(f"⚠  Naming: '{p.stem}' doesn't match {detected_template} convention (expected: {name_example})")
    elif name_example:
        print(f"Naming convention: ✓")
    print(f"\nTo publish: python3 publish.py --file \"{source}\" --space SPACE --parent \"Parent\"")


def main():
    parser = argparse.ArgumentParser(description="Publish documents to Confluence")
    parser.add_argument("--mapping", help="Path to mapping CSV/Excel file")
    parser.add_argument("--file", help="Single file path or Google Doc URL")
    parser.add_argument("--space", help="Space key (e.g. ISMS)")
    parser.add_argument("--parent", help="Parent page title")
    parser.add_argument("--title", help="Page title")
    parser.add_argument("--template", default="general",
                        choices=["policy","procedure","workflow","record","form",
                                 "meeting_minutes","iso27001","general"])
    parser.add_argument("--labels", help="Comma-separated labels")
    parser.add_argument("--dry-run", action="store_true",
                        help="Preview upload plan without publishing")
    parser.add_argument("--preview", action="store_true",
                        help="Render the built ADF as a local HTML page and open in browser before publishing")
    parser.add_argument("--test-auth", action="store_true",
                        help="Test credentials and list spaces, then exit")
    parser.add_argument("--analyze", help="Analyze a file's structure without publishing")
    parser.add_argument("--audit",       metavar="SPACE_KEY",
                        help="Audit all pages in a space for template compliance")
    parser.add_argument("--remediate",   metavar="SPACE_KEY",
                        help="Audit then patch non-compliant pages with placeholder sections")
    parser.add_argument("--folder",      metavar="FOLDER_NAME",
                        help="Limit --audit, --remediate, or --tree to a specific folder")
    parser.add_argument("--go",          action="store_true",
                        help="Skip confirmation prompts and use default publish actions")
    parser.add_argument("--list-spaces",          action="store_true",
                        help="List all Confluence spaces and exit")
    parser.add_argument("--tree",                  metavar="SPACE_KEY",
                        help="Print the full page/folder tree for a space and exit")
    parser.add_argument("--set-regulation",        metavar="REGULATION",
                        help="Set active regulation (e.g. iso27001) and save to .confluence-config.json")
    parser.add_argument("--clear-regulation",      action="store_true",
                        help="Remove the active regulation from .confluence-config.json")
    parser.add_argument("--list-regulation-docs",  metavar="REGULATION",
                        help="List the document catalog for a regulation and exit")
    parser.add_argument("--fix-heading-numbers",   metavar="SPACE_KEY",
                        help="Scan and fix restarting numbered headings in existing Confluence pages")
    parser.add_argument("--add-print-headers",      metavar="SPACE_KEY",
                        help="Add document control header and print footer to pages that are missing them")
    parser.add_argument("--set-policy",            metavar="SOURCE",
                        help="Load formatting rules from a local file or Confluence page URL/ID and save as project style policy")
    parser.add_argument("--policy-section",        metavar="SECTION",
                        help="Extract only a named section from the policy source, e.g. 'Appendix A'")
    args = parser.parse_args()

    # Set regulation mode
    if args.set_regulation:
        key = args.set_regulation.lower().replace('/', '').replace(' ', '')
        if key not in REGULATION_CONFIGS:
            print(f"Unknown regulation '{key}'. Available: {', '.join(REGULATION_CONFIGS.keys())}")
            sys.exit(1)
        cfg = load_regulation_config()
        cfg['regulation']  = key
        cfg['set_date']    = __import__('datetime').date.today().isoformat()
        save_regulation_config(cfg)
        print(f"Active regulation set to: {REGULATION_CONFIGS[key]['name']}")
        print(f"Config saved to: {_REGULATION_CONFIG_FILE}")
        return

    # Clear regulation mode
    if args.clear_regulation:
        cfg = load_regulation_config()
        cfg.pop('regulation', None)
        cfg.pop('set_date', None)
        save_regulation_config(cfg)
        print("Regulation cleared.")
        return

    # List regulation docs
    if args.list_regulation_docs:
        key = args.list_regulation_docs.lower().replace('/', '').replace(' ', '')
        reg = REGULATION_CONFIGS.get(key)
        if not reg:
            print(f"Unknown regulation '{key}'. Available: {', '.join(REGULATION_CONFIGS.keys())}")
            sys.exit(1)
        print(f"\n{reg['name']} — Document Catalog ({len(reg['docs'])} documents)\n")
        print(f"  {'ID':<12}  Document Name")
        print(f"  {'─'*12}  {'─'*50}")
        for doc_id, doc_name in reg['docs'].items():
            print(f"  {doc_id:<12}  {doc_name}")
        print()
        return

    # Fix heading numbers mode
    if args.fix_heading_numbers:
        run_fix_heading_numbers(args.fix_heading_numbers, folder=args.folder, go=args.go)
        return

    # Add print headers mode
    if args.add_print_headers:
        run_add_print_headers(args.add_print_headers, folder=args.folder, go=args.go)
        return

    # Set style policy mode (local files don't need Confluence credentials)
    if args.set_policy:
        run_set_policy(args.set_policy, section=args.policy_section)
        return

    # Analyze mode — no credentials needed
    if args.analyze:
        analyze_file(args.analyze)
        return

    # List spaces mode
    if args.list_spaces:
        if not ATLASSIAN_URL:
            print("ERROR: ATLASSIAN_URL not set in .env")
            sys.exit(1)
        try:
            spaces = list_spaces()
            print(f"\nSpaces ({len(spaces)} total):\n")
            for s in spaces:
                print(f"  {s['key']:<12} — {s['name']:<40}  [{s.get('type','global')}]")
            print()
        except Exception as e:
            print(f"Error fetching spaces: {e}")
            sys.exit(1)
        return

    # Tree mode
    if args.tree:
        if not ATLASSIAN_URL:
            print("ERROR: ATLASSIAN_URL not set in .env")
            sys.exit(1)
        try:
            space_key = args.tree
            print(f"\nBuilding tree for space: {space_key} ...\n")
            if args.folder:
                # Scope tree to a specific folder/page subtree
                parent_id = resolve_parent_id(space_key, args.folder)
                root_pages = walk_descendant_pages(parent_id)
                # Build a minimal flat→tree for the subtree
                nodes_by_id = {p["id"]: {"id": p["id"], "title": p["title"],
                                          "type": p["type"], "children": [],
                                          "parent_id": (p.get("ancestors") or [None])[-1]}
                               for p in root_pages}
                for n in nodes_by_id.values():
                    if isinstance(n["parent_id"], dict):
                        n["parent_id"] = n["parent_id"]["id"]
                roots = []
                for n in nodes_by_id.values():
                    pid = n["parent_id"]
                    if pid and pid in nodes_by_id:
                        nodes_by_id[pid]["children"].append(n)
                    else:
                        roots.append(n)
                roots.sort(key=lambda n: n["title"])
                print(f"📁 {args.folder}")
                _print_tree_nodes(roots, indent=1)
                pages, folders = _count_tree(roots)
                print(f"\n  {pages} page(s), {folders} folder(s) under '{args.folder}'\n")
            else:
                tree = build_space_tree(space_key)
                pages, folders = _count_tree(tree)
                print(f"Space {space_key}  —  {pages} page(s), {folders} folder(s)\n")
                _print_tree_nodes(tree)
                print()
        except Exception as e:
            print(f"Error building tree: {e}")
            sys.exit(1)
        return

    # Audit mode
    if args.audit:
        results = run_audit(args.audit, folder=args.folder)
        print_audit_report(results, args.audit)
        return

    # Remediate mode
    if args.remediate:
        run_remediate(args.remediate, folder=args.folder, go=args.go)
        return

    # Auth test mode
    if args.test_auth:
        if not ATLASSIAN_URL:
            print("ERROR: ATLASSIAN_URL not set in .env")
            sys.exit(1)
        try:
            spaces = list_spaces()
            print(f"Connected! Found {len(spaces)} spaces:")
            for s in spaces:
                print(f"  {s['key']:12} — {s['name']}")
        except Exception as e:
            print(f"Connection failed: {e}")
            sys.exit(1)
        return

    if not ATLASSIAN_URL:
        print("ERROR: ATLASSIAN_URL not set in .env")
        sys.exit(1)

    # Show spaces
    print("\nFetching Confluence spaces...")
    spaces = list_spaces()
    print("\nAvailable spaces:")
    for i, s in enumerate(spaces, 1):
        print(f"  {i}. {s['key']:10} — {s['name']}")
    print()

    if args.dry_run:
        print("[DRY RUN] No pages will be published.\n")

    if args.file:
        _publish_single_file(args)
        return

    if args.mapping:
        _publish_mapping(args)
        return

    parser.print_help()


if __name__ == "__main__":
    main()
